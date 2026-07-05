from __future__ import annotations
from pathlib import Path
from docx import Document


def extract_docx(path: Path) -> list[dict]:
    document = Document(path)
    items: list[dict] = []
    buffer: list[str] = []
    section = "Document body"

    def flush() -> None:
        nonlocal buffer, section
        text = "\n".join(p for p in buffer if p.strip()).strip()
        if text:
            items.append({"text": text, "section": section, "extraction_method": "python-docx"})
        buffer = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "").lower()
        if "heading" in style:
            flush()
            section = text
        else:
            buffer.append(text)
    flush()

    for t_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            items.append({"text": "\n".join(rows), "section": f"Table {t_index}", "extraction_method": "python-docx-table"})
    return items
